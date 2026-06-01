"""文档加载器 - 支持多种文档格式

支持：
- PDF 文档
- Word 文档 (.docx)
- Markdown 文件
- HTML 网页
- 纯文本文件
- JSON 文件
"""

import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional


@dataclass
class Document:
    """文档对象"""

    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    source: str = ""
    page: Optional[int] = None

    def __post_init__(self):
        if not self.metadata:
            self.metadata = {}
        if self.source:
            self.metadata["source"] = self.source
        if self.page is not None:
            self.metadata["page"] = self.page


class DocumentLoader:
    """
    通用文档加载器

    支持多种文档格式的加载和解析
    """

    SUPPORTED_FORMATS = {
        ".txt": "text",
        ".md": "markdown",
        ".markdown": "markdown",
        ".pdf": "pdf",
        ".docx": "docx",
        ".html": "html",
        ".htm": "html",
        ".json": "json",
    }

    def __init__(self, encoding: str = "utf-8"):
        self.encoding = encoding
        self._check_dependencies()

    def _check_dependencies(self):
        """检查依赖包"""
        self._has_pypdf = False
        self._has_docx = False
        self._has_bs4 = False

        try:
            import pypdf

            self._has_pypdf = True
        except ImportError:
            pass

        try:
            import docx

            self._has_docx = True
        except ImportError:
            pass

        try:
            from bs4 import BeautifulSoup

            self._has_bs4 = True
        except ImportError:
            pass

    def load(self, file_path: str) -> List[Document]:
        """
        加载文档

        Args:
            file_path: 文件路径

        Returns:
            文档列表（PDF 等多页文档可能返回多页）
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        suffix = path.suffix.lower()
        loader_type = self.SUPPORTED_FORMATS.get(suffix)

        if not loader_type:
            raise ValueError(f"不支持的文档格式: {suffix}")

        loader_method = getattr(self, f"_load_{loader_type}", None)
        if not loader_method:
            raise ValueError(f"未实现加载器: {loader_type}")

        return loader_method(file_path)

    def load_directory(self, dir_path: str, recursive: bool = False) -> List[Document]:
        """
        加载目录下的所有文档

        Args:
            dir_path: 目录路径
            recursive: 是否递归子目录

        Returns:
            所有文档列表
        """
        path = Path(dir_path)

        if not path.is_dir():
            raise NotADirectoryError(f"不是目录: {dir_path}")

        documents = []

        if recursive:
            pattern = "**/*"
        else:
            pattern = "*"

        for file_path in path.glob(pattern):
            if file_path.is_file():
                suffix = file_path.suffix.lower()
                if suffix in self.SUPPORTED_FORMATS:
                    try:
                        docs = self.load(str(file_path))
                        documents.extend(docs)
                    except Exception as e:
                        print(f"   ⚠️ 加载失败 {file_path}: {e}")

        return documents

    def _load_text(self, file_path: str) -> List[Document]:
        """加载纯文本文件"""
        with open(file_path, "r", encoding=self.encoding) as f:
            content = f.read()

        return [Document(content=content, source=file_path, metadata={"format": "text"})]

    def _load_markdown(self, file_path: str) -> List[Document]:
        """加载 Markdown 文件"""
        with open(file_path, "r", encoding=self.encoding) as f:
            content = f.read()

        # 提取 frontmatter（如果有）
        metadata = {"format": "markdown"}

        if content.startswith("---"):
            parts = content.split("---", 2)
            if len(parts) >= 3:
                # 解析 frontmatter
                fm_content = parts[1].strip()
                for line in fm_content.split("\n"):
                    if ":" in line:
                        key, value = line.split(":", 1)
                        metadata[key.strip()] = value.strip()
                content = parts[2].strip()

        return [Document(content=content, source=file_path, metadata=metadata)]

    def _load_pdf(self, file_path: str) -> List[Document]:
        """加载 PDF 文件"""
        if not self._has_pypdf:
            raise ImportError("pypdf 未安装，请运行: pip install pypdf")

        from pypdf import PdfReader

        reader = PdfReader(file_path)
        documents = []

        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            documents.append(
                Document(
                    content=text,
                    source=file_path,
                    page=i + 1,
                    metadata={"format": "pdf", "total_pages": len(reader.pages)},
                )
            )

        return documents

    def _load_docx(self, file_path: str) -> List[Document]:
        """加载 Word 文档"""
        if not self._has_docx:
            raise ImportError("python-docx 未安装，请运行: pip install python-docx")

        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)

        # 提取段落
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # 提取表格
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                tables_text.append(" | ".join(row_text))

        content = "\n\n".join(paragraphs)
        if tables_text:
            content += "\n\n## 表格\n" + "\n".join(tables_text)

        return [Document(content=content, source=file_path, metadata={"format": "docx"})]

    def _load_html(self, file_path: str) -> List[Document]:
        """加载 HTML 文件"""
        if not self._has_bs4:
            raise ImportError("beautifulsoup4 未安装，请运行: pip install beautifulsoup4")

        from bs4 import BeautifulSoup

        with open(file_path, "r", encoding=self.encoding) as f:
            soup = BeautifulSoup(f, "html.parser")

        # 移除脚本和样式
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()

        # 提取文本
        text = soup.get_text(separator="\n", strip=True)

        # 清理多余空白
        text = re.sub(r"\n\s*\n", "\n\n", text)
        text = re.sub(r"[ \t]+", " ", text)

        # 提取标题作为元数据
        title = soup.title.string if soup.title else ""
        metadata = {"format": "html", "title": title}

        return [Document(content=text, source=file_path, metadata=metadata)]

    def _load_json(self, file_path: str) -> List[Document]:
        """加载 JSON 文件"""
        import json

        with open(file_path, "r", encoding=self.encoding) as f:
            data = json.load(f)

        # 转换为文本
        if isinstance(data, dict):
            content = json.dumps(data, ensure_ascii=False, indent=2)
        elif isinstance(data, list):
            content = json.dumps(data, ensure_ascii=False, indent=2)
        else:
            content = str(data)

        return [Document(content=content, source=file_path, metadata={"format": "json"})]

    def load_from_url(self, url: str) -> List[Document]:
        """
        从 URL 加载网页

        Args:
            url: 网页 URL

        Returns:
            文档列表
        """
        import urllib.request

        with urllib.request.urlopen(url) as response:
            html = response.read().decode(self.encoding)

        # 保存到临时文件并加载
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".html", delete=False) as f:
            f.write(html.encode(self.encoding))
            temp_path = f.name

        try:
            docs = self._load_html(temp_path)
            for doc in docs:
                doc.metadata["url"] = url
            return docs
        finally:
            os.unlink(temp_path)


class FAQLoader:
    """FAQ 文档加载器（专门用于客服知识库）"""

    def load_csv(self, file_path: str) -> List[Document]:
        """加载 CSV 格式的 FAQ"""
        import csv

        documents = []
        with open(file_path, "r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                question = row.get("question", "")
                answer = row.get("answer", "")

                if question and answer:
                    content = f"Q: {question}\n\nA: {answer}"
                    documents.append(
                        Document(
                            content=content,
                            source=file_path,
                            metadata={"format": "faq_csv", "question": question, "answer": answer},
                        )
                    )

        return documents

    def load_json(self, file_path: str) -> List[Document]:
        """加载 JSON 格式的 FAQ"""
        import json

        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        documents = []
        if isinstance(data, list):
            for item in data:
                question = item.get("question", "")
                answer = item.get("answer", "")
                category = item.get("category", "general")

                if question and answer:
                    content = f"Q: {question}\n\nA: {answer}"
                    documents.append(
                        Document(
                            content=content,
                            source=file_path,
                            metadata={
                                "format": "faq_json",
                                "question": question,
                                "answer": answer,
                                "category": category,
                            },
                        )
                    )

        return documents
